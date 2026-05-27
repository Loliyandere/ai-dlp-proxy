"""
dlp/file_extractor.py
---------------------
Extract text từ file đính kèm để DLP scan.

Hỗ trợ:
  - PDF        : pdfplumber
  - DOCX/DOC   : python-docx
  - XLSX/XLS   : openpyxl
  - Text/Code  : utf-8 decode
  - Unknown    : detect từ magic bytes, fallback decode text
"""

import io
import logging
from dataclasses import dataclass
from typing import Optional

logger = logging.getLogger("ai_dlp_proxy.file_extractor")

# Extension → loại xử lý
_TEXT_CODE_EXTS = {
    "txt", "csv", "log",
    "py", "js", "ts", "jsx", "tsx",
    "java", "kt", "go", "rb", "php", "cs", "cpp", "c", "h",
    "sql", "sh", "bash",
    "json", "yaml", "yml", "toml", "ini", "env",
    "xml", "html", "htm", "css",
    "md", "rst", "tf", "dockerfile",
}

# Chỉ skip ảnh, media, archive thực sự
_SKIP_EXTS = {
    "jpg", "jpeg", "png", "gif", "bmp", "svg", "webp", "ico",
    "zip", "rar", "7z", "tar", "gz",
    "exe", "dll", "so",
    "mp3", "mp4", "avi", "mov", "wav",
    # "bin" ĐÃ BỎ → sẽ detect bằng magic bytes
}

# Magic bytes để nhận diện file type từ content
_MAGIC_BYTES = [
    (b"%PDF",             "pdf"),
    (b"PK\x03\x04",       "docx"),   # ZIP-based: docx / xlsx
    (b"\xd0\xcf\x11\xe0", "doc"),    # OLE2: doc / xls cũ
    (b"\xff\xd8\xff",     "jpg"),
    (b"\x89PNG",          "png"),
    (b"GIF8",             "gif"),
    (b"ID3",              "mp3"),
]


def _detect_by_magic(content: bytes) -> Optional[str]:
    """Đoán file type từ magic bytes."""
    for magic, ext in _MAGIC_BYTES:
        if content[:len(magic)] == magic:
            return ext
    # Thử decode UTF-8 → text / source code
    try:
        content[:512].decode("utf-8")
        return "txt"
    except (UnicodeDecodeError, ValueError):
        return None


@dataclass
class ExtractResult:
    text: str
    filename: str
    file_type: str
    char_count: int
    truncated: bool = False


def extract_text(
    content: bytes,
    filename: str,
    max_chars: int = 50_000,
) -> Optional[ExtractResult]:
    """
    Extract text từ file.
    Khi filename không rõ (upload.bin, unknown...) → detect từ magic bytes.
    """
    if not filename:
        filename = "unknown"

    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    # Nếu ext là bin/unknown → detect từ magic bytes
    if ext in ("bin", "unknown", "") or not ext:
        detected = _detect_by_magic(content)
        if detected:
            logger.info(f"[FileExtractor] Detected type by magic bytes: {detected} ({filename})")
            ext = detected
            filename = f"{filename}.{ext}" if "." not in filename else filename
        else:
            logger.debug(f"[FileExtractor] Cannot detect type for {filename}")
            return None

    # Skip ảnh, media
    if ext in _SKIP_EXTS:
        logger.debug(f"[FileExtractor] Skip: {filename}")
        return None

    file_type = ext.upper()
    text = None

    # ── PDF ──────────────────────────────────────────────────────────────────
    if ext == "pdf":
        text = _extract_pdf(content)

    # ── DOCX ─────────────────────────────────────────────────────────────────
    elif ext in ("docx", "doc"):
        text = _extract_docx(content)

    # ── XLSX / XLS ───────────────────────────────────────────────────────────
    elif ext in ("xlsx", "xls"):
        text = _extract_xlsx(content)

    # ── Text / Source code ───────────────────────────────────────────────────
    elif ext in _TEXT_CODE_EXTS:
        text = _extract_plain_text(content)

    # ── Fallback: thử decode text ─────────────────────────────────────────────
    else:
        text = _extract_plain_text(content)
        file_type = "UNKNOWN"

    if not text or not text.strip():
        return None

    text = text.strip()
    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars]

    return ExtractResult(
        text=text,
        filename=filename,
        file_type=file_type,
        char_count=len(text),
        truncated=truncated,
    )


# ── Internal extractors ───────────────────────────────────────────────────────

def _extract_pdf(content: bytes) -> Optional[str]:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        return "\n".join(pages) if pages else None
    except ImportError:
        logger.warning("[FileExtractor] pdfplumber not installed")
        return None
    except Exception as e:
        logger.error(f"[FileExtractor] PDF error: {e}")
        return None


def _extract_docx(content: bytes) -> Optional[str]:
    try:
        from docx import Document
        doc   = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts) if parts else None
    except ImportError:
        logger.warning("[FileExtractor] python-docx not installed")
        return None
    except Exception as e:
        logger.error(f"[FileExtractor] DOCX error: {e}")
        return None


def _extract_xlsx(content: bytes) -> Optional[str]:
    try:
        import openpyxl
        wb   = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows) if rows else None
    except ImportError:
        logger.warning("[FileExtractor] openpyxl not installed")
        return None
    except Exception as e:
        logger.error(f"[FileExtractor] XLSX error: {e}")
        return None


def _extract_plain_text(content: bytes) -> Optional[str]:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return None